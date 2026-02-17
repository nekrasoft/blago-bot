from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".docx", ".doc", ".xlsx", ".pdf"}


class DocumentExtractionError(RuntimeError):
    """Raised when text extraction from document failed."""


class UnsupportedDocumentTypeError(ValueError):
    """Raised when document type is unsupported."""


def _clean_line(line: str) -> str:
    return " ".join(line.replace("\xa0", " ").split()).strip()


def _normalize_text(text: str) -> str:
    lines = (_clean_line(line) for line in text.splitlines())
    return "\n".join(line for line in lines if line)


def extract_document_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(file_path)
    if suffix == ".doc":
        return extract_doc_text(file_path)
    if suffix == ".xlsx":
        return extract_xlsx_text(file_path)
    if suffix == ".pdf":
        return extract_pdf_text(file_path)
    raise UnsupportedDocumentTypeError(
        f"Unsupported document extension: {suffix or '<none>'}. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def extract_docx_text(file_path: Path) -> str:
    doc = Document(str(file_path))
    blocks: list[str] = []

    for paragraph in doc.paragraphs:
        text = _clean_line(paragraph.text)
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_line(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                blocks.append(" | ".join(cells))

    return _normalize_text("\n".join(blocks))


def extract_doc_text(file_path: Path) -> str:
    errors: list[str] = []

    try:
        return _extract_doc_via_libreoffice(file_path)
    except Exception as exc:  # pragma: no cover - system dependency path
        errors.append(f"LibreOffice conversion failed: {exc}")

    for tool in ("antiword", "catdoc"):
        try:
            text = _extract_doc_via_cli_text(tool, file_path)
            if text:
                return text
        except Exception as exc:  # pragma: no cover - system dependency path
            errors.append(f"{tool} failed: {exc}")

    error_details = "; ".join(errors) if errors else "no extraction method available"
    raise DocumentExtractionError(
        "Cannot parse .doc file. Install one of: LibreOffice, antiword, catdoc. "
        f"Details: {error_details}"
    )


def extract_xlsx_text(file_path: Path) -> str:
    workbook = load_workbook(filename=str(file_path), data_only=True, read_only=True)
    blocks: list[str] = []

    for sheet in workbook.worksheets:
        sheet_lines: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [_cell_to_text(value) for value in row]
            values = [value for value in values if value]
            if values:
                sheet_lines.append(" | ".join(values))

        if sheet_lines:
            blocks.append(f"Лист: {sheet.title}")
            blocks.extend(sheet_lines)

    workbook.close()

    text = _normalize_text("\n".join(blocks))
    if not text:
        raise DocumentExtractionError("Cannot parse .xlsx file: workbook is empty")
    return text


def extract_pdf_text(file_path: Path) -> str:
    try:
        reader = PdfReader(str(file_path))
    except Exception as exc:
        raise DocumentExtractionError(f"Cannot open .pdf file: {exc}") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise DocumentExtractionError(f"Cannot decrypt .pdf file: {exc}") from exc

    blocks: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        normalized = _normalize_text(page_text)
        if normalized:
            blocks.append(f"Страница {index}")
            blocks.append(normalized)

    text = _normalize_text("\n".join(blocks))
    if not text:
        raise DocumentExtractionError(
            "Cannot parse .pdf file: no extractable text (possibly scanned document)"
        )
    return text


def _extract_doc_via_libreoffice(file_path: Path) -> str:
    office_binary = _find_binary("soffice", "libreoffice")
    if not office_binary:
        raise FileNotFoundError("LibreOffice binary not found (soffice/libreoffice)")

    with tempfile.TemporaryDirectory(prefix="doc_convert_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        command = [
            office_binary,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(tmp_dir),
            str(file_path),
        ]
        subprocess.run(
            command,
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=90,
        )

        converted = tmp_dir / f"{file_path.stem}.docx"
        if not converted.exists():
            candidates = sorted(tmp_dir.glob("*.docx"))
            if not candidates:
                raise FileNotFoundError("Converted .docx was not created")
            converted = candidates[0]

        text = extract_docx_text(converted)
        if not text:
            raise DocumentExtractionError("Converted .docx is empty")
        return text


def _extract_doc_via_cli_text(tool: str, file_path: Path) -> str:
    binary = shutil.which(tool)
    if not binary:
        raise FileNotFoundError(f"{tool} binary not found")

    result = subprocess.run(
        [binary, str(file_path)],
        check=True,
        capture_output=True,
        timeout=60,
    )
    text = _normalize_text(_decode_text_output(result.stdout))
    if not text:
        raise DocumentExtractionError(f"{tool} returned empty output")
    return text


def _decode_text_output(payload: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def _cell_to_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Да" if value else "Нет"
    return _clean_line(str(value))


def _find_binary(*candidates: str) -> str | None:
    for candidate in candidates:
        found = shutil.which(candidate)
        if found:
            return found
    return None
